# app.py
import streamlit as st
import pandas as pd
import time
import os
from io import BytesIO
from docx import Document
import tempfile
from datetime import datetime
import zipfile
if not os.path.exists("config.py"):
    exec(open("config_page.py", encoding="utf-8").read())
    st.stop()

try:
    import config
except ImportError:
    st.error("No se ha encontrado el archivo de configuración 'config.py'.")
    exec(open("config_page.py", encoding="utf-8").read())
    st.stop()

# Verificamos que config tiene todos los campos necesarios y se agregan los que falten
default_fields = {
    "correo_proveedor": '""',
    "default_template": 'r".default.docx"',
    "nombre": '""',
    "selected_docs": '["Word", "PDF"]',
    "enable_advanced_date_fields": 'True',
    "enable_custom_fields": 'False',
    "enable_different_company": 'True',
    "enable_description": 'True',
    "enable_alcance": 'True',
}

missing_fields = {}
for field, default_value in default_fields.items():
    if not hasattr(config, field):
        missing_fields[field] = default_value

if missing_fields:
    st.error(f"El archivo de configuración no tiene los siguientes campos: {', '.join(missing_fields.keys())}. Se agregan con valores por defecto.")
    with open("config.py", "a", encoding="utf-8") as f:
        for field, default_value in missing_fields.items():
            f.write(f'{field} = {default_value}\n')
    st.stop()
       

st.set_page_config(
    page_title="Creador de Ofertas",
    page_icon="📄"
)

#Guardamos la hora para saber si decir buenos dias, tardes o noches
now  = time.localtime()
greetings = "Buenos días" if now.tm_hour < 12 else "Buenas tardes" if now.tm_hour < 19 else "Buenas noches"

# --- Sidebar (menú lateral) ---
with st.sidebar:
    # Logo clicable que redirige a minsait.com  
    logo_url = "https://pbs.twimg.com/profile_images/1859630278114684929/7BumEThB_200x200.jpg"
    try:
        import requests, base64
        response = requests.get(logo_url)
        response.raise_for_status()
        encoded_string = base64.b64encode(response.content).decode()
        left_col, right_col = st.columns([1, 2])
        with left_col:
            st.markdown(f'<a href="https://minsait.com" target="_blank"><img src="data:image/jpeg;base64,{encoded_string}" width="100" style="margin-top: -20px;"></a>', unsafe_allow_html=True)
        with right_col:
            st.markdown("<h1 style='text-align: left; margin-top: -40px;'>Generador de Ofertas\n</h1>", unsafe_allow_html=True)
            try:
                st.markdown(f"<p style='text-align: left; margin-top: -20px;'>{greetings} {config.nombre}</p>", unsafe_allow_html=True)
            except:
                st.markdown(f"<p style='text-align: left; margin-top: -20px;'>Bienvenido</p>", unsafe_allow_html=True)
    except Exception as e:
        st.warning("Error al cargar")
    st.markdown("---")
    
    # Configuración en expander
    with st.expander("Configuración"):
        # Toggle para habilitar campos personalizados
        
        new_proveedor = st.text_input("Correo Proveedor Predeterminado", value=config.correo_proveedor)
        try:
            plantilla_dir = r".\plantillas"
            template_files = [f for f in os.listdir(plantilla_dir) if f.lower().endswith(".docx")]
        except Exception as e:
            template_files = []
            st.error(f"Error al listar archivos en {plantilla_dir}: {e}")
        new_default_template = st.selectbox(
            "Plantilla por Defecto (Selecciona una)", 
            options=template_files if template_files else [config.default_template],
            index=0,
            help="Reemplaza la plantilla por defecto, recomendado solo si se ha modificado la plantilla original"
        )
        if st.button("📁 Abrir carpeta de plantillas", key="open_default_template_folder"):
            try:
                os.startfile(plantilla_dir) 
            except Exception as e:
                st.error(f"Error al abrir la carpeta de plantillas: {e}")


        selected_docs = st.multiselect("Documentos a generar", options=["Word", "PDF"], default=config.selected_docs, help="Selecciona los documentos a generar")
        enable_advanced_date_fields = st.toggle("Habilitar campos de fecha avanzados",
                                                value=getattr(config, "enable_advanced_date_fields", True), help="Permite seleccionar fechas con calendario")      
        enable_different_company = st.toggle("Habilitar diferentes empresas",
                value=getattr(config, "enable_different_company", True), help="Permite seleccionar diferentes empresas para el cliente y proveedor")
        enable_custom_fields = st.toggle("Habilitar campos personalizados", 
                value=getattr(config, "enable_custom_fields", False), help="Permite añadir campos personalizados en el documento")
        enable_description = st.toggle("Habilitar descripción",
                  value=getattr(config, "enable_description", True), help="Permite añadir descripción en el documento")
        enable_alcance = st.toggle("Habilitar alcance",
                  value=getattr(config, "enable_alcance", True), help="Permite añadir alcance en el documento")
        
        if st.button("Guardar Configuración"):
            with open("config.py", "w", encoding="utf-8") as f:
                f.write(f'correo_proveedor = "{new_proveedor}"\n')
                f.write(f'default_template = r"{new_default_template}"\n')
                f.write(f'nombre = "{config.nombre}"\n')
                f.write(f'selected_docs = {selected_docs}\n')
                f.write(f'enable_advanced_date_fields = {enable_advanced_date_fields}\n')
                f.write(f'enable_different_company = {enable_different_company}\n')
                f.write(f'enable_custom_fields = {enable_custom_fields}\n')
                f.write(f'enable_description = {enable_description}\n')
                f.write(f'enable_alcance = {enable_alcance}\n')
                st.success("Configuración guardada. Reinicia la app para aplicar cambios.")
        
    st.markdown("---")
    st.markdown("**Version:** 1.5.3")

# --- Carga de archivos ---
col1, col2= st.columns(2)
with col1:
    excel_file = st.file_uploader("Selecciona el archivo Excel (.xlsx)", type=["xlsx"])
    # Boton para descargar plantilla vacia
    excel_template_path = r".\plantilla.xlsx"
    if os.path.exists(excel_template_path):
        with open(excel_template_path, "rb") as f:
            st.download_button(
                label="Descargar plantilla Excel vacía",
                data=f,
                file_name="oferta.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
    else:
        st.warning("Plantilla Excel no encontrada")
    
with col2:
    template_file = st.file_uploader("Selecciona la plantilla Word (.docx)", type=["docx"])
    # Boton para descargar plantilla vacia

    word_template_path = os.path.join("plantillas", config.default_template) if os.path.exists(os.path.join("plantillas", config.default_template)) else config.default_template
    if os.path.exists(word_template_path):
        with open(word_template_path, "rb") as f:
            st.download_button(
                label="Descargar plantilla Word vacía",
                data=f,
                file_name=config.default_template,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("Plantilla Word por defecto no encontrada")
    

# --- Campos de texto para ubicación de salida y plantilla por defecto ---
st.text_input("Plantilla predeterminada (Se usará en caso de no seleccionar otra plantilla):", value=config.default_template, disabled=True)

if not excel_file:
    st.warning("Para continuar por favor, sube la Plantilla POST.")
    st.stop()

# --- Función para extraer datos del Excel ---
def extraer_datos_excel(excel_file):
    df = pd.read_excel(excel_file, header=None, sheet_name='Plantilla POST')
    oferta_referencia = excel_file.name.split('.')[0]
    try:
        nombre_proyecto = df.loc[6, 1] if pd.notna(df.loc[6, 1]) else None
    except:
        nombre_proyecto = None
    try:
        razon_social_proveedor = df.loc[2, 1] if pd.notna(df.loc[2, 1]) else None
    except:
        razon_social_proveedor = None
    try:
        cif_proveedor = df.loc[3, 1] if pd.notna(df.loc[3, 1]) else None
    except:
        cif_proveedor = None
    try:
        gps_proveedor = df.loc[4, 1] if pd.notna(df.loc[4, 1]) else None
    except:
        gps_proveedor = None
    try:
        fecha_inicio = pd.to_datetime(df.loc[3, 6], format='%d.%m.%Y').strftime('%d/%m/%Y')
    except:
        fecha_inicio = None
    try:
        fecha_fin = pd.to_datetime(df.loc[4, 6], format='%d.%m.%Y').strftime('%d/%m/%Y')
    except:
        fecha_fin = None
    # Extraer SDA de la celda B8 (fila 8, columna 2: índices 7,1)
    try:
        sda = df.loc[7, 1] if pd.notna(df.loc[7, 1]) else None
    except:
        sda = ""
    
    # Extracción de posts (máx. 10)
    posts = []
    for i in range(10, 78):
        if df.loc[i, 6] != 0:
            precio = '{:,.2f}'.format(float(df.loc[i, 6])).replace(',', 'X').replace('.', ',').replace('X', '.')
            posts.append({"post": df.loc[i, 0], "horas": df.loc[i, 3], "costo": precio})
            if len(posts) >= 10:
                break

    # Totales
    try:
        totalh = df.iloc[78, 3] if pd.notna(df.iloc[78, 3]) else None
    except:
        totalh = None
    try:
        totalsiva = '{:,.2f}'.format(float(df.loc[6, 6])).replace(',', 'X').replace('.', ',').replace('X', '.') if pd.notna(df.loc[6, 6]) else None
    except:
        totalsiva = None
    try:
        totalciva = '{:,.2f}'.format(float(df.loc[7, 6])).replace(',', 'X').replace('.', ',').replace('X', '.') if pd.notna(df.loc[7, 6]) else None
    except:
        totalciva = None
    today = time.strftime("%d/%m/%Y")
    
    return {
        "oferta_referencia": oferta_referencia,
        "razon_social_proveedor": razon_social_proveedor,
        "cif_proveedor": cif_proveedor,
        "gps_proveedor": gps_proveedor,
        "nombre_proyecto": nombre_proyecto,
        "fecha_inicio": fecha_inicio,
        "fecha_fin": fecha_fin,
        "descripcion": "",
        "alcance": "",
        "correo_cliente": "",
        "correo_proveedor": config.correo_proveedor,
        "sda": sda,
        "posts": posts,
        "totalh": totalh,
        "totalsiva": totalsiva,
        "totalciva": totalciva,
        "today": today
    }

data = extraer_datos_excel(excel_file)

# --- Control de cantidad de posts en session_state ---
if "n_posts" not in st.session_state:
    st.session_state.n_posts = len(data["posts"]) if len(data["posts"]) > 0 else 1

# =============================================================================
# SECCIÓN 1: DATOS GENERALES
# =============================================================================
with st.container():
    st.subheader("Datos Generales", anchor=None)
    with st.form("form_datos_generales", clear_on_submit=False):
        oferta_referencia = st.text_input("Oferta de Referencia", value=data["oferta_referencia"], placeholder="<<oferta_referencia>>", help="Este campo se toma automaticamente segun el nombre del archivo Excel")
        # Permitir editar el nombre original del proyecto
        nombre_proyecto_original = st.text_input("Nombre del Proyecto", value=data["nombre_proyecto"], help="Celda B7", placeholder = "<<nombre_proyecto>>")
        # Nuevo campo SDA (editable), auto-rellenado desde celda B8
        sda_field = st.text_input("SDA", value=data["sda"], placeholder = "Si hay SDA, saldrá despues del nombre del proyecto entre parentesis", help="Celda B8")
        # Actualiza el nombre del proyecto según SDA: si SDA tiene contenido, concatena; si no, solo el nombre original
        if sda_field is None:
            nombre_proyecto = nombre_proyecto_original
        else:
            nombre_proyecto = f"{nombre_proyecto_original} ({sda_field.strip()})"

        if config.enable_different_company:
            # Hay 3 campos, Razon Social Proveedor(B3), CIF Proveedor(B4) y y Nº de GPS del Proveedor(B5)
            razon_social_proveedor_original = st.text_input("Razón Social Proveedor", value=data["razon_social_proveedor"], help="Celda B3", placeholder = "<<razon_social_proveedor>>")
            cif_proveedor = st.text_input("CIF Proveedor", value=data["cif_proveedor"], help="Celda B4", placeholder = "<<cif_proveedor>>")
            gps_proveedor = st.text_input("Nº de GPS Proveedor", value=data["gps_proveedor"], help="Celda B5", placeholder = "<<gps_proveedor>>")
            if gps_proveedor is not None:
                razon_social_proveedor = f"{razon_social_proveedor_original} ({gps_proveedor})"
            else:
                razon_social_proveedor = razon_social_proveedor_original


        #----Fechas----#
        if config.enable_advanced_date_fields:
            if data["fecha_inicio"] is None:
                fecha_inicio = st.date_input("Fecha de Inicio", value=None, format="DD/MM/YYYY", help="Celda G4, <<fecha_inicio>>")
            else:
                fecha_inicio = st.date_input("Fecha de Inicio", value=datetime.strptime(data["fecha_inicio"], '%d/%m/%Y').date(), format="DD/MM/YYYY", help="Celda G4, <<fecha_inicio>>").strftime('%d/%m/%Y')
            if data["fecha_fin"] is None:
                fecha_fin = st.date_input("Fecha de Fin", value=None, format="DD/MM/YYYY", help="Celda G5, <<fecha_fin>>")
            else:
                fecha_fin = st.date_input("Fecha de Fin", value=datetime.strptime(data["fecha_fin"], '%d/%m/%Y').date(), format="DD/MM/YYYY", help="Celda G5, <<fecha_fin>>").strftime('%d/%m/%Y')
        else:
            fecha_inicio = st.text_input("Fecha de Inicio", value=data["fecha_inicio"], help="Celda G4", placeholder = "<<fecha_inicio>>")
            fecha_fin = st.text_input("Fecha de Fin", value=data["fecha_fin"], help="Celda G5", placeholder = "<<fecha_fin>>")
        #--------------#

        correo_cliente = st.text_input("Correo Cliente", value=data["correo_cliente"], help="No data", placeholder = "<<correo_cliente>>")
        correo_proveedor = st.text_input("Correo Proveedor", value=data["correo_proveedor"], help = "Configurable - No data", placeholder = "<<correo_proveedor>>")
        if config.enable_advanced_date_fields:
            today = st.date_input("Today", value="today", format = "DD/MM/YYYY", help="Fecha de hoy")
        else:
            today = st.text_input("Today", value=data["today"], help="Fecha de hoy", placeholder = "<<today>>")
        # Descripción y Alcance
        if config.enable_description:
            descripcion = st.text_area("Descripción", value=data["descripcion"], help = "Ignorar si la plantilla ya contiene la descripción",  placeholder = "<<descripcion>>")
        else:
            descripcion = ""
        if config.enable_alcance:
            alcance = st.text_area("Alcance", value=data["alcance"], help = "Ignorar si la plantilla ya contiene el alcance", placeholder = "<<alcance>>")
        else:
            alcance = ""
        #--------------#

        submitted_dg = st.form_submit_button("Guardar Datos Generales")
        if submitted_dg:
            st.success("Datos Generales guardados.")

        # Ponemos las fechas en formato dd/mm/yyyy
        try:
            fecha_inicio = fecha_inicio.strftime('%d/%m/%Y')
        except:
            pass
        try:
            fecha_fin = fecha_fin.strftime('%d/%m/%Y')
        except:
            pass
        try:
            today = today.strftime('%d/%m/%Y')
        except:
            pass

# =============================================================================
# SECCIÓN 2: POSTS Y TOTALES
# =============================================================================
with st.container():
    st.subheader("Perfiles", anchor=None)
    col_action1, col_action2 = st.columns(2)
    with col_action1:
        if st.button("Agregar Perfil", key="agregar_post", help="Agrega un Perfil POST hasta un maximo de 10") and st.session_state.n_posts < 10:
            st.session_state.n_posts += 1
    with col_action2:
        if st.button("Borrar Perfil", key="borrar_post", help = "Borrar un Perfil POST") and st.session_state.n_posts > 1:
            st.session_state.n_posts -= 1

    with st.form("form_posts", clear_on_submit=False):
        st.markdown("#### Edición de Perfiles Posts")
        posts = []
        for i in range(st.session_state.n_posts):
            if i < len(data["posts"]):
                default_post = data["posts"][i]
            else:
                default_post = {"post": "", "horas": "", "costo": ""}
            col1, col2, col3 = st.columns([2, 1, 1])
            with col1:
                post_val = st.text_input(f"Perfil Post {i+1}", value=default_post["post"], key=f"post_{i}", placeholder=f"<<post{i+1}>>", help=f"Columna A11:A78")
            with col2:
                post_horas = st.text_input(f"Horas Perfil Post {i+1}", value=default_post["horas"], key=f"posth_{i}", placeholder=f"<<posth{i+1}>>", help=f"Columna D11:D78")
            with col3:
                post_costo = st.text_input(f"Costo Perfil Post {i+1}", value=default_post["costo"], key=f"postc_{i}", placeholder=f"<<postc{i+1}>>", help=f"Columna G11:G78")
            posts.append({"post": post_val, "horas": post_horas, "costo": post_costo})
        
        st.markdown("#### Totales")
        colT1, colT2 = st.columns(2)
        with colT1:
            totalh = st.text_input("Total Horas", value=str(data["totalh"]), placeholder="<<totalh>>", help="Celda D79")
        with colT2:
            totalsiva = st.text_input("Total sin IVA", value=str(data["totalsiva"]), placeholder="<<totalsiva>>", help="Celda G7")
            totalciva = st.text_input("Total con IVA", value=str(data["totalciva"]), placeholder="<<totalciva>>", help="Celda G8")
        submitted_posts = st.form_submit_button("Guardar Posts y Totales")
    if submitted_posts:
        st.success("Posts y Totales guardados.")

# =============================================================================
# SECCIÓN 3: CAMPOS PERSONALIZADOS(Opcional)
# =============================================================================
if config.enable_custom_fields:
    with st.container():
        st.subheader("Campos Personalizados", anchor=None)
        
        # Inicializar contadores en 0 para que empiece sin ningun campo
        if "n_custom_fields" not in st.session_state:
            st.session_state.n_custom_fields = 0
        if "n_large_custom_fields" not in st.session_state:
            st.session_state.n_large_custom_fields = 0

        # Botones para agregar y quitar campos:
        # Agregar Campo, Quitar Campo, Agregar Campo Grande, Quitar Campo Grande
        col_add, col_remove, col_add_large, col_remove_large = st.columns(4)
        with col_add:
            if st.button("Agregar Campo"):
                st.session_state.n_custom_fields += 1
        with col_remove:
            if st.button("Quitar Campo") and st.session_state.n_custom_fields > 0:
                st.session_state.n_custom_fields -= 1
        with col_add_large:
            if st.button("Agregar Parrafo"):
                st.session_state.n_large_custom_fields += 1
        with col_remove_large:
            if st.button("Quitar Parrafo") and st.session_state.n_large_custom_fields > 0:
                st.session_state.n_large_custom_fields -= 1

        # Formulario para campos personalizados
        with st.form("form_campos_personalizados", clear_on_submit=False):
            custom_fields = []
            
            # Verificar si hay campos de texto cortos o grandes
            if st.session_state.n_custom_fields == 0 and st.session_state.n_large_custom_fields == 0:
                st.markdown("#### No hay ningun campo")
            else:
                if st.session_state.n_custom_fields > 0:
                    st.markdown("#### Campos de Texto Cortos")
                    for i in range(st.session_state.n_custom_fields):
                        col_placeholder, col_reemplazo = st.columns(2)
                        with col_placeholder:
                            placeholder_val = st.text_input(f"Placeholder Campo {i+1}", value=f"<<campo{i+1}>>", key=f"ph_campo_{i}")
                        with col_reemplazo:
                            replacement_val = st.text_input(f"Reemplazo Campo {i+1}", value="", key=f"campo_{i}")
                        custom_fields.append({"placeholder": placeholder_val, "replacement": replacement_val})
                if st.session_state.n_large_custom_fields > 0:
                    st.markdown("#### Campos de Parrafos")
                    for j in range(st.session_state.n_large_custom_fields):
                        col_placeholder, col_reemplazo = st.columns(2)
                        with col_placeholder:
                            placeholder_val = st.text_input(f"Placeholder Parrafos {j+1}", value=f"<<campo_grande{j+1}>>", key=f"ph_campo_grande_{j}")
                        with col_reemplazo:
                            replacement_val = st.text_area(f"Reemplazo Parrafo {j+1}", value="", key=f"campo_grande_{j}")
                        custom_fields.append({"placeholder": placeholder_val, "replacement": replacement_val})
            
            submitted_cp = st.form_submit_button("Guardar Campos Personalizados")
            if submitted_cp:
                st.success("Campos Personalizados guardados.")

# =============================================================================
# SECCIÓN FINAL: GENERAR DOCUMENTO
# =============================================================================
with st.container():
    if st.button("Generar Documento", help="Genera documento(s) según la configuración seleccionada"):
        updated = {
            "oferta_referencia": oferta_referencia,
            "razon_social_proveedor": razon_social_proveedor if config.enable_different_company else "",
            "cif_proveedor": cif_proveedor if config.enable_different_company else "",
            "nombre_proyecto": nombre_proyecto,
            "fecha_inicio": fecha_inicio,
            "fecha_fin": fecha_fin,
            "correo_cliente": correo_cliente,
            "correo_proveedor": correo_proveedor,
            "descripcion": descripcion,
            "alcance": alcance,
            "today": today,
            "posts": posts,
            "totalh": totalh,
            "totalsiva": totalsiva,
            "totalciva": totalciva
        }
        
        st.subheader("Datos Actualizados")
        df_generales = pd.DataFrame({
            "Campo": ["Oferta de Referencia", "Nombre del Proyecto", "Fecha de Inicio", "Fecha de Fin", 
                      "Correo Cliente", "Correo Proveedor", "Today"],
            "Valor": [updated["oferta_referencia"], updated["nombre_proyecto"], updated["fecha_inicio"], 
                      updated["fecha_fin"], updated["correo_cliente"], updated["correo_proveedor"], 
                      updated["today"]]
        })
        if config.enable_different_company:
            if config.enable_different_company:
                custom_df = pd.DataFrame({
                     "Campo": ["Razón Social Proveedor", "CIF Proveedor"],
                     "Valor": [updated["razon_social_proveedor"], updated["cif_proveedor"]]
                })
                df_generales = pd.concat([custom_df, df_generales], ignore_index=True)
                # Reordenamos(indice 2 al indice 0)
                df_generales = pd.concat([df_generales.iloc[[2]], df_generales.iloc[:2], df_generales.iloc[3:]], ignore_index=True)
        if config.enable_description:
            df_generales = pd.concat([df_generales, pd.DataFrame({"Campo": ["Descripción"], "Valor": [updated["descripcion"]]})], ignore_index=True)
        if config.enable_alcance:
            df_generales = pd.concat([df_generales, pd.DataFrame({"Campo": ["Alcance"], "Valor": [updated["alcance"]]})], ignore_index=True)


        st.table(df_generales)
        
        if updated["posts"]:
            df_posts = pd.DataFrame(updated["posts"])
            st.subheader("Posts")
            st.table(df_posts)
        
        totales_df = pd.DataFrame({
            "Descripción": ["Total Horas", "Total sin IVA", "Total con IVA"],
            "Valor": [updated["totalh"], updated["totalsiva"], updated["totalciva"]]
        })
        st.subheader("Totales")
        st.table(totales_df)
        
        progress_bar = st.progress(0)
        
        # Cargar la plantilla Word
        if template_file is None:
            st.info(f"No se cargó plantilla; se usará la plantilla por defecto:\n {config.default_template}")
            doc = Document(os.path.join("plantillas", config.default_template))
        else:
            template_file.seek(0)
            doc = Document(template_file)
        progress_bar.progress(5)
        
        # Reemplazo de placeholders generales y totales
        placeholders = {
            "<<oferta_referencia>>": updated["oferta_referencia"],
            "<<proveedor>>": updated["razon_social_proveedor"],
            "<<cif>>": updated["cif_proveedor"],           
            "<<razon_social_proveedor>>": updated["razon_social_proveedor"],
            "<<cif_proveedor>>": updated["cif_proveedor"],
            "<<nombre_proyecto>>": updated["nombre_proyecto"],
            "<<fecha_inicio>>": updated["fecha_inicio"],
            "<<fecha_fin>>": updated["fecha_fin"],
            "<<correo_cliente>>": updated["correo_cliente"],
            "<<correo_proveedor>>": updated["correo_proveedor"],
            "<<descripcion>>": updated["descripcion"],
            "<<alcance>>": updated["alcance"],
            "<<totalh>>": updated["totalh"],
            "<<totalsiva>>": updated["totalsiva"],
            "<<totalciva>>": updated["totalciva"],
            "<<today>>": updated["today"]
        }
        for ph, val in placeholders.items():
            for p in doc.paragraphs:
                if ph in p.text:
                    p.text = p.text.replace(ph, str(val))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if ph in cell.text:
                            cell.text = cell.text.replace(ph, str(val))
        progress_bar.progress(15)
        
        # Reemplazo de placeholders para cada post
        for i, post in enumerate(updated["posts"], start=1):
            ph_post = f"<<post{i}>>"
            ph_posth = f"<<posth{i}>>"
            ph_postc = f"<<postc{i}>>"
            for p in doc.paragraphs:
                if ph_post in p.text:
                    p.text = p.text.replace(ph_post, post["post"])
                if ph_posth in p.text:
                    p.text = p.text.replace(ph_posth, post["horas"])
                if ph_postc in p.text:
                    p.text = p.text.replace(ph_postc, post["costo"])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if ph_post in cell.text:
                            cell.text = cell.text.replace(ph_post, post["post"])
                        if ph_posth in cell.text:
                            cell.text = cell.text.replace(ph_posth, post["horas"])
                        if ph_postc in cell.text:
                            cell.text = cell.text.replace(ph_postc, post["costo"])
        progress_bar.progress(50)
        
        # Procesamiento de la tabla de posts: borrar filas vacías y formatear algunas celdas
        try:
            table = doc.tables[1]
            for row in table.rows[-2:]:
                for cell in row.cells:
                    if cell.paragraphs and len(cell.paragraphs[0].runs) > 0:
                        cell.paragraphs[0].runs[0].bold = True
            rows_to_delete = []
            for i, row in enumerate(table.rows):
                cell_text = row.cells[0].text.strip()
                if cell_text == "-" or (cell_text.startswith("<<") and cell_text.endswith(">>")):
                    rows_to_delete.append(i)
            for i in sorted(rows_to_delete, reverse=True):
                row = table.rows[i]
                row._element.getparent().remove(row._element)
        except Exception as e:
            st.error(f"Error al procesar la tabla de posts: {e}")
        
        # Reemplazo de campos personalizados
        if getattr(config, "enable_custom_fields", False):
            for field in custom_fields:
                for p in doc.paragraphs:
                    if field["placeholder"] in p.text:
                        p.text = p.text.replace(field["placeholder"], field["replacement"])
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if field["placeholder"] in cell.text:
                                cell.text = cell.text.replace(field["placeholder"], field["replacement"])
        progress_bar.progress(55)
        
        # Formateo extra en el documento Word
        try:
            if nombre_proyecto is not None:
                doc.paragraphs[1].runs[0].bold = True
                para = doc.paragraphs[5]
                text = para.runs[0].text
                para.runs[0].text = ""
                run1 = para.add_run(text[:19])
                run1.underline = True
                run2 = para.add_run(text[19:])
                run2.underline = False
        except Exception as e:
            st.error(f"Error en formateo extra del documento: {e}")

        # Formateo de Razon Social Proveedor y CIF Proveedor
        try:
            if config.enable_different_company:
                # Subrayar PROVEEDOR: <<proveedor>>
                para_proveedor = doc.paragraphs[2]
                for run in para_proveedor.runs:
                    run.underline = True

                # Subrayar CIF: <<cif>>
                para_cif = doc.paragraphs[3]
                for run in para_cif.runs:
                    run.underline = True
        except Exception as e:
            st.error(f"Error al subrayar: {e}")






        progress_bar.progress(60)
        
        doc_filename = f"{updated['oferta_referencia']}.docx"
        pdf_filename = f"{updated['oferta_referencia']}.pdf"
        
        # Decisión basada en config.selected_docs: si se selecciona un único documento o ambos
        if len(config.selected_docs) == 1 or len(config.selected_docs) == 0:
            if "Word" in config.selected_docs or len(config.selected_docs) == 0:
                if len(config.selected_docs) == 0:
                    st.warning("No se ha seleccionado un documento de descarga específico. Por defecto se generará un documento Word.")
                # Genera únicamente documento Word (descarga)
                word_io = BytesIO()
                doc.save(word_io)
                word_io.seek(0)
                progress_bar.progress(100)
                st.download_button(
                    "Descargar documento Word",
                    data=word_io,
                    file_name=doc_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_word_single"
                )
            elif "PDF" in config.selected_docs:
                # Aviso de conversión a PDF
                st.info("Convirtiendo a PDF, por favor espere...")
                # Genera únicamente documento PDF (conversión desde el Word generado) (descarga)
                if not os.path.exists(r".\temp"):
                    os.makedirs(r".\temp")
                with tempfile.TemporaryDirectory(dir=r".\temp") as temp_folder:
                    temp_doc_path = os.path.join(temp_folder, "temp.docx")
                    doc.save(temp_doc_path)
                    try:
                        import pythoncom
                        pythoncom.CoInitialize()
                        from docx2pdf import convert
                        temp_pdf_path = os.path.join(temp_folder, pdf_filename)
                        convert(temp_doc_path, temp_pdf_path)
                        with open(temp_pdf_path, "rb") as f_pdf:
                            pdf_bytes = f_pdf.read()
                        pdf_io = BytesIO(pdf_bytes)
                        pdf_io.seek(0)
                        progress_bar.progress(100)
                        st.download_button(
                            "Descargar documento PDF",
                            data=pdf_io,
                            file_name=pdf_filename,
                            mime="application/pdf",
                            key="download_pdf_single"
                        )
                    except Exception as e:
                        st.error(f"Error al convertir a PDF: {e}")
        elif len(config.selected_docs) >= 2:
            # Genera ambos documentos y los empaqueta en un ZIP (descarga)
            temp_doc_path = r".\temp\temp.docx"
            word_io = BytesIO()
            doc.save(word_io)
            word_io.seek(0)
            
            if not os.path.exists(r".\temp"):
                os.makedirs(r".\temp")
            with tempfile.TemporaryDirectory(dir=r".\temp") as temp_folder:
                try:
                    # Aviso de conversión a PDF
                    st.info("Convirtiendo a PDF, por favor espere...")
                    import pythoncom
                    pythoncom.CoInitialize()
                    from docx2pdf import convert
                    temp_pdf_path = os.path.join(temp_folder, pdf_filename)
                    # Guardar el documento Word temporalmente en la carpeta .\temp
                    temp_doc_fullpath = os.path.join(temp_folder, "temp.docx")
                    doc.save(temp_doc_fullpath)
                    convert(temp_doc_fullpath, temp_pdf_path)
                    with open(temp_pdf_path, "rb") as pdf_file:
                        pdf_bytes = pdf_file.read()
                    pdf_io = BytesIO(pdf_bytes)
                    pdf_io.seek(0)
                    progress_bar.progress(100)
                except Exception as e:
                    st.error(f"Error al convertir a PDF: {e}")
                    pdf_io = None
                    st.error(f"Error al convertir a PDF: {e}")
                    pdf_io = None

            excel_file.seek(0)
            excel_bytes = excel_file.read()
            zip_io = BytesIO()
            with zipfile.ZipFile(zip_io, mode="w", compression=zipfile.ZIP_DEFLATED) as zip_file:
                zip_file.writestr(doc_filename, word_io.getvalue())
                if pdf_io is not None:
                    zip_file.writestr(pdf_filename, pdf_io.getvalue())
                zip_file.writestr(excel_file.name, excel_bytes)
            zip_io.seek(0)
            st.download_button(
                "Descargar ZIP con todos los archivos",
                data=zip_io,
                file_name=f"{updated['oferta_referencia']}.zip",
                mime="application/zip"
            )
