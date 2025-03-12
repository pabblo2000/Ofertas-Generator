import streamlit as st
import pandas as pd
import time
import os
from io import BytesIO
from docx import Document
import config  # Configuración: correo_proveedor, output_folder, default_template

# --- Sidebar (menú lateral) ---
with st.sidebar:
    # Logo clicable que redirige a minsait.com  
    logo_path = "https://pbs.twimg.com/profile_images/1859630278114684929/7BumEThB_200x200.jpg"
    try:
        import requests, base64
        response = requests.get(logo_path)
        response.raise_for_status()
        encoded_string = base64.b64encode(response.content).decode()
        st.markdown(f'<a href="https://minsait.com" target="_blank" ><img src="data:image/jpeg;base64,{encoded_string}" width="150"></a>', unsafe_allow_html=True)
        st.markdown("<h1 style='text-align: left; color: #FF5733;'>Generador de Ofertas</h1>", unsafe_allow_html=True)
    except Exception as e:
        st.warning("No se pudo cargar el logo desde la URL.")
    
    st.markdown("---")
    
    # Configuración en expander
    with st.expander("Configuración"):
        new_proveedor = st.text_input("Correo Proveedor Configuración", value=config.correo_proveedor)
        new_output_folder = st.text_input("Ruta de Salida", value=config.output_folder)
        new_default_template = st.text_input("Plantilla por Defecto", value=config.default_template)
        if st.button("Guardar Configuración"):
            with open("config.py", "w") as f:
                f.write(f'correo_proveedor = "{new_proveedor}"\n')
                f.write(f'output_folder = r"{new_output_folder}"\n')
                f.write(f'default_template = r"{new_default_template}"\n')
            st.success("Configuración guardada. Reinicia la app para aplicar cambios.")

    st.markdown("---")
    st.markdown("**Version:** 1.00")
    st.markdown("**Autor:** Pablo Álvaro")

# --- Carga de archivos ---
col1, col2 = st.columns(2)
with col1:
    excel_file = st.file_uploader("Selecciona el archivo Excel (.xlsx)", type=["xlsx"])
with col2:
    template_file = st.file_uploader("Selecciona la plantilla Word (.docx)", type=["docx"])

if not excel_file:
    st.warning("Por favor, sube el archivo Excel.")
    st.stop()

# --- Función para extraer datos del Excel ---
def extraer_datos_excel(excel_file):
    df = pd.read_excel(excel_file, header=None, sheet_name='Plantilla POST')
    oferta_referencia = excel_file.name.split('.')[0]
    nombre_proyecto = df.loc[6, 1]
    fecha_inicio = pd.to_datetime(df.loc[3, 6], format='%d.%m.%Y').strftime('%d/%m/%Y')
    fecha_fin = pd.to_datetime(df.loc[4, 6], format='%d.%m.%Y').strftime('%d/%m/%Y')
    sda = df.loc[7, 1]
    if pd.notna(sda):
        nombre_proyecto = f"{nombre_proyecto} ({sda})"
    
    # Extracción de posts (máx. 5)
    posts = []
    for i in range(10, 78):
        if df.loc[i, 6] != 0:
            precio = '{:,.2f}'.format(float(df.loc[i, 6])).replace(',', 'X').replace('.', ',').replace('X', '.')
            posts.append({"post": df.loc[i, 0], "horas": df.loc[i, 3], "costo": precio})
            if len(posts) >= 5:
                break

    # Totales
    totalh = df.iloc[78, 3]
    totalsiva = '{:,.2f}'.format(float(df.loc[6, 6])).replace(',', 'X').replace('.', ',').replace('X', '.')
    totalciva = '{:,.2f}'.format(float(df.loc[7, 6])).replace(',', 'X').replace('.', ',').replace('X', '.')
    today = time.strftime("%d/%m/%Y")
    
    return {
        "oferta_referencia": oferta_referencia,
        "nombre_proyecto": nombre_proyecto,
        "fecha_inicio": fecha_inicio,
        "fecha_fin": fecha_fin,
        "descripcion": "",
        "correo_cliente": "",
        "correo_proveedor": config.correo_proveedor,  # Valor predeterminado desde la configuración
        "posts": posts,
        "totalh": totalh,
        "totalsiva": totalsiva,
        "totalciva": totalciva,
        "today": today
    }

data = extraer_datos_excel(excel_file)

# --- Control de cantidad de posts en session_state ---
if "n_posts" not in st.session_state:
    st.session_state.n_posts = len(data["posts"]) if len(data["posts"]) > 0 else 1

# =============================================================================
# SECCIÓN 1: DATOS GENERALES
# =============================================================================
with st.container():
    st.subheader("Datos Generales")
    with st.form("form_datos_generales", clear_on_submit=False):
        oferta_referencia = st.text_input("Oferta de Referencia", value=data["oferta_referencia"])
        nombre_proyecto = st.text_input("Nombre del Proyecto", value=data["nombre_proyecto"])
        fecha_inicio = st.text_input("Fecha de Inicio", value=data["fecha_inicio"])
        fecha_fin = st.text_input("Fecha de Fin", value=data["fecha_fin"])
        correo_cliente = st.text_input("Correo Cliente", value=data["correo_cliente"])
        correo_proveedor = st.text_input("Correo Proveedor", value=data["correo_proveedor"])
        today = st.text_input("Today", value=data["today"])
        descripcion = st.text_area("Descripción", value=data["descripcion"])
        submitted_dg = st.form_submit_button("Guardar Datos Generales")
    if submitted_dg:
        st.success("Datos Generales guardados.")

# =============================================================================
# SECCIÓN 2: POSTS Y TOTALES
# =============================================================================
with st.container():
    st.subheader("Posts")
    if st.button("Agregar Post", key="agregar_post") and st.session_state.n_posts < 5:
        st.session_state.n_posts += 1

    with st.form("form_posts", clear_on_submit=False):
        st.markdown("### Edición de Posts")
        posts = []
        for i in range(st.session_state.n_posts):
            if i < len(data["posts"]):
                default_post = data["posts"][i]
            else:
                default_post = {"post": "", "horas": "", "costo": ""}
            # Ajustamos el ancho: 2 para Post, 1 para Horas y 1 para Costo
            col1, col2, col3 = st.columns([2, 1, 1])
            with col1:
                post_val = st.text_input(f"Post {i+1}", value=default_post["post"], key=f"post_{i}")
            with col2:
                post_horas = st.text_input(f"Horas Post {i+1}", value=default_post["horas"], key=f"posth_{i}")
            with col3:
                post_costo = st.text_input(f"Costo Post {i+1}", value=default_post["costo"], key=f"postc_{i}")
            posts.append({"post": post_val, "horas": post_horas, "costo": post_costo})
        
        st.markdown("### Totales")
        colT1, colT2 = st.columns(2)
        with colT1:
            totalh = st.text_input("Total Horas", value=str(data["totalh"]))
        with colT2:
            totalsiva = st.text_input("Total sin IVA", value=str(data["totalsiva"]))
            totalciva = st.text_input("Total con IVA", value=str(data["totalciva"]))
        submitted_posts = st.form_submit_button("Guardar Posts y Totales")
    if submitted_posts:
        st.success("Posts y Totales guardados.")

# =============================================================================
# SECCIÓN FINAL: GENERAR DOCUMENTO
# =============================================================================
if st.button("Generar Documento"):
    updated = {
        "oferta_referencia": oferta_referencia,
        "nombre_proyecto": nombre_proyecto,
        "fecha_inicio": fecha_inicio,
        "fecha_fin": fecha_fin,
        "correo_cliente": correo_cliente,
        "correo_proveedor": correo_proveedor,
        "descripcion": descripcion,
        "today": today,
        "posts": posts,
        "totalh": totalh,
        "totalsiva": totalsiva,
        "totalciva": totalciva
    }
    
    st.subheader("Datos Actualizados")
    df_generales = pd.DataFrame({
        "Campo": ["Oferta de Referencia", "Nombre del Proyecto", "Fecha de Inicio", "Fecha de Fin", "Correo Cliente", "Correo Proveedor", "Descripción", "Today"],
        "Valor": [updated["oferta_referencia"], updated["nombre_proyecto"], updated["fecha_inicio"], updated["fecha_fin"], updated["correo_cliente"], updated["correo_proveedor"], updated["descripcion"], updated["today"]]
    })
    st.table(df_generales)
    
    if updated["posts"]:
        df_posts = pd.DataFrame(updated["posts"])
        st.subheader("Posts")
        st.table(df_posts)
    
    totales_df = pd.DataFrame({
        "Descripción": ["Total Horas", "Total sin IVA", "Total con IVA"],
        "Valor": [updated["totalh"], updated["totalsiva"], updated["totalciva"]]
    })
    st.subheader("Totales")
    st.table(totales_df)
    
    # --- Generación del documento Word y conversión a PDF con barra de progreso ---
    progress_bar = st.progress(0)
    # Seleccionar plantilla: si no se carga, se usa la por defecto
    if template_file is None:
        st.info(f"No se cargó plantilla; se usará la plantilla por defecto:\n{config.default_template}")
        doc = Document(config.default_template)
    else:
        template_file.seek(0)
        doc = Document(template_file)
    progress_bar.progress(20)
    
    # Reemplazo de placeholders en datos generales y totales
    placeholders = {
        "<<oferta_referencia>>": updated["oferta_referencia"],
        "<<nombre_proyecto>>": updated["nombre_proyecto"],
        "<<fecha_inicio>>": updated["fecha_inicio"],
        "<<fecha_fin>>": updated["fecha_fin"],
        "<<correo_cliente>>": updated["correo_cliente"],
        "<<correo_proveedor>>": updated["correo_proveedor"],
        "<<descripcion>>": updated["descripcion"],
        "<<totalh>>": updated["totalh"],
        "<<totalsiva>>": updated["totalsiva"],
        "<<totalciva>>": updated["totalciva"],
        "<<today>>": updated["today"]
    }
    for ph, val in placeholders.items():
        for p in doc.paragraphs:
            if ph in p.text:
                p.text = p.text.replace(ph, str(val))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if ph in cell.text:
                        cell.text = cell.text.replace(ph, str(val))
    progress_bar.progress(50)
    
    # Reemplazo de placeholders para cada post (<<post1>>, <<posth1>>, <<postc1>>, etc.)
    for i, post in enumerate(updated["posts"], start=1):
        ph_post = f"<<post{i}>>"
        ph_posth = f"<<posth{i}>>"
        ph_postc = f"<<postc{i}>>"
        for p in doc.paragraphs:
            if ph_post in p.text:
                p.text = p.text.replace(ph_post, post["post"])
            if ph_posth in p.text:
                p.text = p.text.replace(ph_posth, post["horas"])
            if ph_postc in p.text:
                p.text = p.text.replace(ph_postc, post["costo"])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if ph_post in cell.text:
                        cell.text = cell.text.replace(ph_post, post["post"])
                    if ph_posth in cell.text:
                        cell.text = cell.text.replace(ph_posth, post["horas"])
                    if ph_postc in cell.text:
                        cell.text = cell.text.replace(ph_postc, post["costo"])
    progress_bar.progress(80)
    
    # --- Borrar filas vacías de la tabla de perfiles POST y formatear ---
    try:
        table = doc.tables[1]
        
        # Poner en negrita las últimas dos filas
        for row in table.rows[-2:]:
            for cell in row.cells:
                if cell.paragraphs and len(cell.paragraphs[0].runs) > 0:
                    cell.paragraphs[0].runs[0].bold = True

        # Borrar las filas que no estén rellenadas
        rows_to_delete = []
        for i, row in enumerate(table.rows):
            cell_text = row.cells[0].text.strip()
            # Se borra si es "-" o si parece un placeholder (ej: <<post1>>)
            if cell_text == "-" or (cell_text.startswith("<<") and cell_text.endswith(">>")):
                rows_to_delete.append(i)
        
        for i in sorted(rows_to_delete, reverse=True):
            row = table.rows[i]
            row._element.getparent().remove(row._element)
    except Exception as e:
        st.error(f"Error al procesar la tabla de posts: {e}")
    
    # --- Formateo extra en el documento Word ---
    try:
        # Ponemos la oferta de referencia en negrita (asumiendo que es el párrafo 1)
        doc.paragraphs[1].runs[0].bold = True
        
        # Subrayar únicamente las primeras 19 letras del párrafo 5
        para = doc.paragraphs[5]
        text = para.runs[0].text
        para.runs[0].text = ""
        run1 = para.add_run(text[:19])
        run1.underline = True
        run2 = para.add_run(text[19:])
        run2.underline = False
    except Exception as e:
        st.error(f"Error en formateo extra del documento: {e}")
    progress_bar.progress(90)
    
    # Guardar el documento Word y convertirlo a PDF en la carpeta configurada
    doc_filename = f"{updated['oferta_referencia']}_generado.docx"
    pdf_filename = f"{updated['oferta_referencia']}_generado.pdf"
    doc_path = os.path.join(config.output_folder, doc_filename)
    doc.save(doc_path)
    progress_bar.progress(95)
    try:
        from docx2pdf import convert
        pdf_path = os.path.join(config.output_folder, pdf_filename)
        convert(doc_path, pdf_path)
        progress_bar.progress(100)
        st.success(f"Documentos generados correctamente en:\n{config.output_folder}")
    except Exception as e:
        st.error(f"El documento Word se generó en {config.output_folder}, pero hubo un error al convertir a PDF: {e}")
